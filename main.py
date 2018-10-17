from fuzzywuzzy import process
from tqdm import tqdm
import fire
import docx
import os
import csv

def docx2csv(**kwargs):
    for k_, v_ in kwargs.items():
        setattr(options, k_, v_)

    for i in range(1, 51):
        doc = docx.Document(f'./docs/天龍八部_大意{str(i).zfill(2)}.docx')
        table = doc.tables[0]
        rows = table.rows
        with open(f'./csv/天龍八部_大意{str(i).zfill(2)}.csv', 'w') as f:
            for row in rows:
                texts = []
                cells = list(row.cells)
                cells.pop(0)
                for cell in cells:
                    texts.append(cell.paragraphs[0].text.replace('\n', ''))

                texts = ','.join(texts)
                f.write(texts + '\n')

def pair_gen(**kwargs):
    for k_, v_ in kwargs.items():
        setattr(options, k_, v_)

    PAIR = open('pair.csv', 'a')
    ERROR = open('error.txt', 'a')
    data = open('sky_dragon_traditional.txt').read().replace('\n', '<newline>')
    array = data.split('<newline>')

    PAIR.write('summary,text\n')
    for i in tqdm(range(1,51)):
        with open(f'./csv/天龍八部_大意{str(i).zfill(2)}.csv') as f:
            reader = csv.reader(f, delimiter=',')
            
            for idx, row in enumerate(tqdm(reader)):
                if idx != 0:
                    begin = row[0]
                    end = row[1]
                    summary = row[2]
                
                    begin = process.extractOne(begin, tqdm(array))[0]
                    end = process.extractOne(end, tqdm(array))[0]
                    text = ''.join(list(data)[data.find(begin) : data.find(end) + len(end)])
                    if text == '':
                        ERROR.write(f'FILE{i} 編號{idx} 空字串: ' + ','.join(row))
                    
                    PAIR.write(f'{summary},{text}\n')

if __name__ == '__main__':
    fire.Fire()