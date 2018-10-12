import docx
import os

doc = docx.Document(f'./docs/天龍八部_大意01.docx')
table = doc.tables[0]

rows = table.rows

for i in range(1, 51):
    with open(f'./csv/天龍八部_大意{str(i).zfill(2)}.csv', 'w') as f:
        for row in rows:
            texts = []
            cells = list(row.cells)
            cells.pop(0)
            for cell in cells:
                texts.append(cell.paragraphs[0].text.replace('\n', ''))

            texts = ','.join(texts)
            f.write(texts + '\n')